from docx import Document

# Создаем новый документ Word
doc = Document()
doc.add_heading('Задание: Проверка палиндрома и обратный порядок строки (Python)', 0)

# Добавляем описание задания
doc.add_paragraph(
    '''Задание: Напишите программу, которая будет проверять, является ли введенная строка палиндромом.
    Строка считается палиндромом, если она читается одинаково слева направо и справа налево, при этом игнорируются
    пробелы и знаки препинания. Например, "А роза упала на лапу Азора" является палиндромом.
    
    Дополнительно, если строка является палиндромом, программа должна вывести эту строку в обратном порядке.
    
    Требования:
    1. Программа должна игнорировать регистр символов.
    2. Программа должна игнорировать пробелы и знаки препинания.
    3. Программа должна вывести строку в обратном порядке, если она является палиндромом.
    4. В случае, если строка не является палиндромом, выведите соответствующее сообщение.
    '''
)

# Добавляем код Python
doc.add_heading('Решение на Python:', level=1)

doc.add_paragraph('''
import re

def is_palindrome(s):
    # Очищаем строку: приводим к нижнему регистру и убираем все, кроме букв
    cleaned_str = re.sub(r'[^a-zA-Z]', '', s.lower())
    return cleaned_str == cleaned_str[::-1]

def reverse_string(s):
    # Функция для вывода строки в обратном порядке
    return s[::-1]

# Ввод от пользователя
user_input = input("Введите строку: ")

# Проверяем, является ли строка палиндромом
if is_palindrome(user_input):
    print("Это палиндром.")
    print(f"Строка в обратном порядке: {reverse_string(user_input)}")
else:
    print("Это не палиндром.")
''')

# Сохраняем документ
file_path = 'Palindrome_Check_Python_Task.docx'
doc.save(file_path)

file_path  # Возвращаем путь к файлу для скачивания



# TODO: Заметки
## Автор: Дуплей Максим Игоревич
## Дата: 08.12.2024
## @quadd4rv1n7 @dupley_maxim_1999