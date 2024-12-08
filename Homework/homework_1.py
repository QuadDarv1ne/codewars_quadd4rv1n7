from docx import Document

# Создаем новый документ Word
doc = Document()
doc.add_heading('Задание: Игра "Горячо - холодно" (Python)', 0)

# Добавляем описание задания
doc.add_paragraph(
    '''Правила игры:
1. Я загадываю число от 1 до 100.
2. Ваша цель — угадать это число за как можно меньшее количество попыток.
3. После каждой попытки я буду говорить, насколько близко вы подошли к загаданному числу:
   - Если ваше число слишком далеко — я скажу "Холодно".
   - Если вы близки, но не угадали — я скажу "Тепло".
   - Если вы почти угадали — я скажу "Горячо".
   - Когда вы угадаете число — я скажу "Поздравляю, вы угадали!"

Начало игры:
- Пишите ваши предположения, а я буду сообщать, насколько близко вы к правильному ответу.
'''
)

# Добавляем код Python
doc.add_heading('Решение на Python:', level=1)

doc.add_paragraph('''
import random

def get_user_guess():
    """Запросить у пользователя число и вернуть его."""
    while True:
        try:
            user_guess = int(input("Введите ваше предположение: "))
            if 1 <= user_guess <= 100:  # Проверяем, что число в допустимом диапазоне
                return user_guess
            else:
                print("Число должно быть от 1 до 100.")
        except ValueError:
            print("Пожалуйста, введите целое число.")

def play_game():
    # Загадать случайное число от 1 до 100
    secret_number = random.randint(1, 100)
    attempts = 0

    print("Я загадал число от 1 до 100. Попробуйте угадать.")

    while True:
        # Запрашиваем число у пользователя
        user_guess = get_user_guess()
        attempts += 1

        # Проверка на совпадение
        if user_guess == secret_number:
            print(f"Поздравляю :D Вы угадали число {secret_number} за {attempts} попыток.")
            break

        # Проверка близости
        difference = abs(user_guess - secret_number)
        if difference <= 5:
            print("Горячо")
        elif difference <= 10:
            print("Тепло")
        else:
            print("Холодно")

# Запуск игры
if __name__ == "__main__":
    play_game()
''')

# Сохраняем документ
file_path = 'Game - Hot&Cold.docx'
doc.save(file_path)

file_path  # Возвращаем путь к файлу для скачивания



# TODO: Заметки
## Автор: Дуплей Максим Игоревич
## Дата: 08.12.2024
## @quadd4rv1n7 @dupley_maxim_1999