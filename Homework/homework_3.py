from docx import Document

# Задание 6
doc1 = Document()
doc1.add_heading('Задание 6: Четные и нечетные числа (Python)', 0)

# Описание задания
doc1.add_paragraph(
    '''Задание: Программа запрашивает у пользователя список целых чисел, разделённых пробелами,
    и делит его на два списка: один с четными числами, другой — с нечетными. Затем оба списка выводятся на экран.

    Требования:
    1. Программа должна корректно обрабатывать список чисел.
    2. Программа должна разделять числа на четные и нечетные.
    3. Результат работы программы должен выводить два списка на экран.

    Пример ввода:
    Введите числа через пробел: 1 2 3 4 5 6 7 8 9 10

    Пример вывода:
    Четные числа: [2, 4, 6, 8, 10]
    Нечетные числа: [1, 3, 5, 7, 9]
    '''
)

# Код решения
doc1.add_heading('Решение на Python:', level=1)

doc1.add_paragraph('''
# Программа для разделения чисел на четные и нечетные

# Ввод чисел от пользователя
numbers = input("Введите числа через пробел: ").split()

# Преобразуем в список целых чисел
numbers = [int(num) for num in numbers]

# Разделение на четные и нечетные числа
even_numbers = [num for num in numbers if num % 2 == 0]
odd_numbers = [num for num in numbers if num % 2 != 0]

# Вывод результатов
print("Четные числа:", even_numbers)
print("Нечетные числа:", odd_numbers)
''')

# Сохраняем документы
file_path1 = 'Even_Odd_Numbers_Task_Python.docx'

doc1.save(file_path1)

file_path1 # Возвращаем пути для скачивания



# TODO: Заметки
## Автор: Дуплей Максим Игоревич
## Дата: 08.12.2024
## @quadd4rv1n7 @dupley_maxim_1999
